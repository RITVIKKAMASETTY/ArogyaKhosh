import os
import traceback
import logging
from celery import shared_task
from django.conf import settings
from django.db import transaction
from web3 import Web3
from eth_account import Account
from eth_account.signers.local import LocalAccount
from web3.exceptions import ContractLogicError
import requests
from .hash import *
from .models import patient, patientDocument, DocumentProcessStatus

# Langchain imports
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.schema import SystemMessage, HumanMessage
from langchain.text_splitter import CharacterTextSplitter
from io import BytesIO

# Set up logger
logger = logging.getLogger(__name__)
GROQ_API_KEY = settings.GROQ_API_KEY# Should be in environment variables

# Initialize or load vector store for medical knowledge
def get_medical_knowledge_vectorstore():
    """Initialize or load medical knowledge vector store for RAG"""
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        knowledge_base_path = os.path.join(settings.BASE_DIR, "medical_knowledge_base")
        
        # Check if vector store exists
        if os.path.exists(os.path.join(knowledge_base_path, "index.faiss")):
            logger.info("Loading existing medical knowledge vector store")
            return FAISS.load_local(knowledge_base_path, embeddings)
        else:
            logger.info("Creating new medical knowledge vector store")
            # Sample medical knowledge - replace with actual medical reference documents
            medical_knowledge = [
                "Medical reports often contain vital signs including blood pressure, heart rate, temperature, and respiratory rate.",
                "Laboratory tests such as CBC, metabolic panels, lipid profiles, and urinalysis provide critical diagnostic information.",
                "Imaging reports from X-rays, CT scans, MRIs, and ultrasounds should be summarized with key findings.",
                "Surgical history should include procedure type, date, complications, and outcomes.",
                "Medication lists should include drug names, dosages, frequencies, and the conditions they treat.",
                "Patient allergies must be prominently noted, including the type of reaction.",
                "Family history should highlight conditions with genetic components or risk factors.",
                "Medical summaries should organize information by body systems or chronologically.",
                "Chronic conditions require tracking of disease progression, treatment responses, and complications.",
                "Recent hospitalizations should note admission reason, treatments, length of stay, and discharge disposition."
            ]
            
            # Convert to Langchain document format
            from langchain.schema import Document
            docs = [Document(page_content=text) for text in medical_knowledge]
            
            # Create vector store
            vector_store = FAISS.from_documents(docs, embeddings)
            
            # Save for future use
            os.makedirs(knowledge_base_path, exist_ok=True)
            vector_store.save_local(knowledge_base_path)
            
            return vector_store
    except Exception as e:
        logger.error(f"Error initializing medical knowledge vector store: {str(e)}")
        return None

# Try to initialize the vector store at module level
try:
    medical_vectorstore = get_medical_knowledge_vectorstore()
    logger.info("Medical knowledge vector store initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize medical knowledge vector store: {str(e)}")
    medical_vectorstore = None

@shared_task(bind=True)
def upload_document_to_ipfs_and_blockchain(self, patient_id, file_name, file_content, summary=False):
    """
    Async task to upload document to IPFS and blockchain
    
    Args:
        patient_id (int): ID of the patient
        file_name (str): Name of the file
        file_content: File object or content
        summary (bool): Whether to update patient summary with file content
    
    Returns:
        dict: Result of the upload process
    """
    # Create process status entry
    process_status = DocumentProcessStatus.objects.create(
        patient=patient.objects.get(id=patient_id),
        file_name=file_name,
        status='PROCESSING',
        task_id=self.request.id
    )
    
    # Track if summary was actually updated
    summary_actually_updated = False

    try:
        # Fetch patient
        try:
            pat = patient.objects.get(id=patient_id)
            logger.info(f"Patient found: {pat.id}")
        except patient.DoesNotExist:
            process_status.status = 'FAILED'
            process_status.error_message = 'Patient not found'
            process_status.save()
            return {'error': 'Patient not found'}

        # Process summary if required
        if summary:
            logger.info(f"Starting summary processing for file: {file_name}")
            try:
                # Extract text from file_content based on file type
                file_text = None
                
                # Check if file_content is a file-like object or bytes
                if hasattr(file_content, 'read'):
                    # It's a file-like object
                    logger.info("Processing file-like object")
                    file_content_bytes = file_content.read()
                    file_content.seek(0)  # Reset file pointer for later use
                else:
                    # It's already bytes
                    logger.info("Processing bytes content")
                    file_content_bytes = file_content
                
                logger.info(f"File content type: {type(file_content)}, Size: {len(file_content_bytes) if isinstance(file_content_bytes, bytes) else 'unknown'}")
                
                # Check if it's a PDF file by examining content (more reliable than extension)
                is_pdf = file_content_bytes[:5] == b'%PDF-' or file_name.lower().endswith('.pdf')
                
                if is_pdf:
                    logger.info("Detected PDF content, using PyPDF2 for extraction")
                    try:
                        import PyPDF2
                        
                        pdf_file = BytesIO(file_content_bytes)
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        file_text = ""
                        logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text() or ""
                            file_text += page_text + "\n"
                            logger.info(f"Extracted {len(page_text)} chars from page {page_num+1}")
                        
                        logger.info(f"Total extracted text from PDF: {len(file_text)} characters")
                        
                    except ImportError:
                        logger.error("PyPDF2 not installed, attempting basic extraction")
                        process_status.error_message = 'Warning: PyPDF2 not installed, cannot extract PDF properly'
                        process_status.save()
                        # Don't fall back to basic extraction for PDFs - it won't work
                        file_text = "Unable to extract PDF content properly. Please install PyPDF2."
                    except Exception as e:
                        logger.error(f"Error extracting PDF text: {str(e)}")
                        process_status.error_message = f'Warning: PDF extraction error: {str(e)}'
                        process_status.save()
                        file_text = f"Error extracting PDF content: {str(e)}"
                elif file_name.lower().endswith(('.txt', '.csv', '.md')):
                    # Simple text files
                    logger.info("Extracting text from text file")
                    file_text = file_content_bytes.decode('utf-8', errors='replace')
                elif file_name.lower().endswith(('.docx')):
                    # For DOCX files
                    logger.info("Extracting text from DOCX file")
                    try:
                        import docx
                        
                        doc = docx.Document(BytesIO(file_content_bytes))
                        file_text = "\n".join([para.text for para in doc.paragraphs])
                        logger.info(f"Extracted {len(file_text)} characters from DOCX")
                    except ImportError:
                        logger.error("python-docx not installed, cannot extract text from DOCX")
                        process_status.error_message = 'Warning: python-docx not installed, cannot extract text from DOCX'
                        process_status.save()
                    except Exception as e:
                        logger.error(f"Error extracting DOCX text: {str(e)}")
                        process_status.error_message = f'Warning: DOCX extraction error: {str(e)}'
                        process_status.save()
                else:
                    # Try basic UTF-8 decoding for unknown types but check for binary content
                    logger.info(f"Attempting basic text extraction for {file_name}")
                    try:
                        sample = file_content_bytes[:1000]
                        non_printable = sum(1 for b in sample if b < 32 and b not in (9, 10, 13))  # Tab, LF, CR
                        
                        # If more than 10% non-printable, likely binary
                        if non_printable > len(sample) * 0.1:
                            logger.warning("File appears to be binary, not text")
                            file_text = "Unable to extract text from binary file"
                        else:
                            file_text = file_content_bytes.decode('utf-8', errors='replace')
                            logger.info(f"Extracted {len(file_text)} characters using basic UTF-8 decoding")
                    except Exception as e:
                        logger.error(f"Error in basic text extraction: {str(e)}")
                        process_status.error_message = f'Warning: Unable to extract text from {file_name}: {str(e)}'
                        process_status.save()
                
                if not file_text or len(file_text.strip()) == 0:
                    logger.error("No text could be extracted from the document")
                    process_status.error_message = 'Warning: Could not extract text for summary generation or extracted text is empty'
                    process_status.save()
                else:
                    # Get existing summary
                    existing_summary = pat.summary or ""
                    logger.info(f"Existing summary length: {len(existing_summary)}")
                    
                    # For debugging purposes, log first few characters of extracted text
                    logger.info(f"Extracted text sample (first 200 chars): {file_text[:200]}")
                    
                    # Truncate text if too large to avoid Groq API limit
                    # Using a rough estimate: 1 token ≈ 4 characters for English text
                    max_chars = 4000  # Leave room for system prompt and existing summary (~1000 tokens)
                    
                    if len(file_text) > max_chars:
                        logger.warning(f"Extracted text too large ({len(file_text)} chars), truncating to {max_chars} chars")
                        file_text = file_text[:max_chars] + "\n\n[Text truncated due to size limitations]"
                    
                    # Instead of directly calling Groq API, use Langchain with RAG
                    logger.info("Starting RAG-based summary generation")
                    
                    try:
                        # Prepare query for RAG
                        query = f"Create a medical summary for: {file_text[:200]}..."
                        
                        # If we have a vector store, retrieve relevant medical knowledge
                        retrieved_context = ""
                        if medical_vectorstore:
                            logger.info("Retrieving relevant medical knowledge")
                            retriever = medical_vectorstore.as_retriever(search_kwargs={"k": 3})
                            relevant_docs = retriever.get_relevant_documents(query)
                            
                            retrieved_context = "\n".join([doc.page_content for doc in relevant_docs])
                            logger.info(f"Retrieved context: {retrieved_context[:200]}...")
                        
                        # Initialize the Groq chat model through Langchain
                        groq_api_key = settings.GROQ_API_KEY
                        
                        if not groq_api_key:
                            logger.error("GROQ_API_KEY not set in settings")
                            process_status.error_message = 'Warning: GROQ_API_KEY not set'
                            process_status.save()
                            return {'error': 'GROQ_API_KEY not set'}
                        
                        chat_model = ChatGroq(
                            api_key=groq_api_key,
                            model_name="llama3-70b-8192",
                            temperature=0.3,
                            max_tokens=1024
                        )
                        
                        # Create system message with retrieved context
                        context_prefix = ""
                        if retrieved_context:
                            context_prefix = f"""Consider the following medical knowledge when creating your summary:
                            {retrieved_context}
                            
                            """
                        
                        system_message = SystemMessage(
                            content=f"""{context_prefix}You are a medical documentation assistant. Create concise medical summaries in Markdown format.
                            Never include patient names or personal identifiers. Do not include recommendations or lifestyle advice unless explicitly stated in the source documents.
                            Format the summary with appropriate Markdown headers (# for main title, ## for sections, ### for subsections),
                            bullet points (- for items), and text formatting (**bold** for important values and diagnoses).
                            Always include sections for: Previous Summary, New Findings, Diagnosis, and any other relevant clinical sections.
                            Focus exclusively on medical findings, test results, and diagnoses. Include the previous summary information.
                            The final summary must be properly formatted in Markdown."""
)

                        
                        # Create user message
                        user_message = HumanMessage(
                            content=f"Existing summary:\n{existing_summary}\n\nNew medical report:\n{file_text}\n\nCreate an updated clinical summary with only the essential medical information."
                        )
                        
                        # Generate response
                        logger.info("Calling Groq model through Langchain")
                        response = chat_model.invoke([system_message, user_message])
                        
                        # Extract the response content
                        updated_summary = response.content
                        
                        logger.info(f"Updated summary received, length: {len(updated_summary)}")
                        
                        # Update patient summary
                        pat.summary = updated_summary
                        pat.save()
                        logger.info("Patient summary saved successfully")
                        summary_actually_updated = True
                        
                    except Exception as e:
                        logger.error(f"Error in RAG-based summary generation: {str(e)}")
                        logger.error(traceback.format_exc())
                        process_status.error_message = f'Warning: RAG summary generation error: {str(e)}, falling back to direct API call'
                        process_status.save()
                        
                        # Fall back to direct Groq API call if RAG fails
                        logger.info("Falling back to direct Groq API call")
                        groq_api_url = "https://api.groq.com/openai/v1/chat/completions"
                        groq_api_key = settings.GROQ_API_KEY
                        
                        groq_headers = {
                            "Authorization": f"Bearer {groq_api_key}",
                            "Content-Type": "application/json"
                        }
                        
                        payload = {
                            "model": "llama3-70b-8192",
                            "messages": [
                                {
                                    "role": "system",
                                    "content": "You are a medical documentation assistant. Create concise medical summaries in Markdown format. Never include patient names or personal identifiers. Format the summary with appropriate Markdown headers (# for main title, ## for sections, ### for subsections), bullet points (- for items), and text formatting (**bold** for important values and diagnoses). Always include sections for: Previous Summary, New Findings, Diagnosis, and any other relevant clinical sections. Focus exclusively on medical findings, test results, and diagnoses. Include the previous summary information. The final summary must be properly formatted in Markdown."
                                },
                                {
                                    "role": "user",
                                    "content": f"Existing summary:\n{existing_summary}\n\nNew medical report:\n{file_text}\n\nCreate an updated clinical summary with the essential medical information formatted in Markdown."
                                }
                            ],
                            "temperature": 0.3
                        }
                                                
                        try:
                            groq_response = requests.post(groq_api_url, headers=groq_headers, json=payload, timeout=30)
                            
                            logger.info(f"Groq API response status: {groq_response.status_code}")
                            if groq_response.status_code == 200:
                                result = groq_response.json()
                                updated_summary = result["choices"][0]["message"]["content"]
                                
                                logger.info(f"Updated summary received, length: {len(updated_summary)}")
                                
                                # Update patient summary
                                pat.summary = updated_summary
                                pat.save()
                                logger.info("Patient summary saved successfully")
                                summary_actually_updated = True
                            else:
                                logger.error(f"Groq API error: {groq_response.text}")
                                process_status.error_message = f'Warning: Summary update failed: {groq_response.text}, but continuing with upload'
                                process_status.save()
                        except requests.exceptions.RequestException as e:
                            logger.error(f"Request error when calling Groq API: {str(e)}")
                            process_status.error_message = f'Warning: Groq API request error: {str(e)}'
                            process_status.save()
            
            except Exception as e:
                # Don't fail the whole process if summary fails
                logger.error(f"Summary processing error: {str(e)}")
                logger.error(traceback.format_exc())
                process_status.error_message = f'Warning: Summary processing error: {str(e)}, but continuing with upload'
                process_status.save()

        # Upload to IPFS
        try:
            logger.info(f"Starting IPFS upload for file: {file_name}")
            pinata_url = "https://api.pinata.cloud/pinning/pinFileToIPFS"
            pinata_headers = {
                "pinata_api_key": settings.PINATA_API_KEY,
                "pinata_secret_api_key": settings.PINATA_SECRET_KEY,
            }
            
            # Ensure file_content is prepared properly for upload
            if hasattr(file_content, 'read'):
                # It's a file-like object, can be used directly in files dictionary
                files = {"file": (file_name, file_content)}
                logger.info("Using file-like object for Pinata upload")
            else:
                # It's bytes, wrap it for the request
                files = {"file": (file_name, BytesIO(file_content))}
                logger.info("Using BytesIO wrapper for Pinata upload")

            pinata_response = requests.post(pinata_url, headers=pinata_headers, files=files)

            if pinata_response.status_code != 200:
                process_status.status = 'FAILED'
                process_status.error_message = f'IPFS upload failed: {pinata_response.text}'
                process_status.save()
                logger.error(f"IPFS upload failed: {pinata_response.text}")
                return {'error': f'Failed to upload to IPFS: {pinata_response.text}'}
                
            ipfs_data = pinata_response.json()
            cid = ipfs_data["IpfsHash"]
            logger.info(f"Received url: {cid}")
            # Optionally encrypt or hash the CID
            hashcid = encrypt_url(cid)  # Assuming this function exists
            logger.info(f"Encrypted url: {hashcid}")
        except Exception as e:
            process_status.status = 'FAILED'
            process_status.error_message = f'IPFS upload error: {str(e)}'
            process_status.save()
            logger.error(f"IPFS upload error: {str(e)}")
            logger.error(traceback.format_exc())
            return {'error': f'IPFS upload error: {str(e)}'}

        # Blockchain transaction
        try:
            logger.info("Starting blockchain transaction")
            # Web3 connection setup
            import json

            web3 = Web3(Web3.HTTPProvider(settings.SEPOLIA_NODE_URL))
            with open(r'/home/sathwik/ArogyaKhosh/backend/core/home/abi.json', "r") as abi_file:
                contract_abi = json.load(abi_file)
            # Account setup
            account: LocalAccount = Account.from_key(settings.ETHEREUM_PRIVATE_KEY)
            sender_address = account.address
            
            # Contract setup
            contract_address = Web3.to_checksum_address(settings.CONTRACT_ADDRESS)
            contract = web3.eth.contract(address=contract_address, abi=contract_abi)
            
            # Create patient document
            patd = patientDocument.objects.create(
                name=file_name,
                patient=pat,
                hash=cid,
            )
            logger.info(f"Created patient document: {patd.id}")
            
            # Get blockchain details
            current_timestamp = int(web3.eth.get_block('latest').timestamp)
            nonce = web3.eth.get_transaction_count(sender_address)
            gas_price = web3.eth.gas_price
            
            # Estimate and build transaction
            estimated_gas = contract.functions.addPatientDocument(
                pat.id, patd.id, str(hashcid), current_timestamp, summary
            ).estimate_gas({'from': sender_address})
            
            gas_limit = int(estimated_gas * 1.2)
            logger.info(f"Estimated gas: {estimated_gas}, Gas limit: {gas_limit}")
            
            transaction = contract.functions.addPatientDocument(
                pat.id, patd.id, str(hashcid), current_timestamp, summary
            ).build_transaction({
                'from': sender_address,
                'nonce': nonce,
                'gas': gas_limit,
                'gasPrice': gas_price,
                'chainId': 11155111  # Sepolia chain ID
            })
            
            # Sign and send transaction
            signed_txn = account.sign_transaction(transaction)
            tx_hash = web3.eth.send_raw_transaction(signed_txn.raw_transaction)
            logger.info(f"Transaction sent: {tx_hash.hex()}")
            
            # Wait for receipt
            receipt = web3.eth.wait_for_transaction_receipt(tx_hash)
            logger.info(f"Transaction receipt status: {receipt.status}")
            
            # Update process status
            process_status.status = 'COMPLETED' if receipt.status else 'FAILED'
            process_status.save()
            
            result = {
                "message": "File uploaded to IPFS and stored on Blockchain",
                "cid": cid,
                "url": f"https://gateway.pinata.cloud/ipfs/{cid}",
                "transaction": tx_hash.hex()
            }
            
            # Add summary info if summary was actually updated
            if summary_actually_updated:
                result["summary_updated"] = True
            else:
                # If summary was requested but failed, include this in the result
                if summary:
                    result["summary_requested"] = True
                    result["summary_updated"] = False
                    result["summary_error"] = process_status.error_message
            
            return result
        
        except Exception as e:
            process_status.status = 'FAILED'
            process_status.error_message = f'Blockchain error: {str(e)}'
            process_status.save()
            logger.error(f"Blockchain error: {str(e)}")
            logger.error(traceback.format_exc())
            return {'error': f'Blockchain error: {str(e)}'}
    
    except Exception as e:
        process_status.status = 'FAILED'
        process_status.error_message = f'Unexpected error: {str(e)}'
        process_status.save()
        logger.error(f"Unexpected error: {str(e)}")
        logger.error(traceback.format_exc())
        return {'error': f'Unexpected error: {str(e)}'}